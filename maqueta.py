from docx.enum.text import WD_ALIGN_PARAGRAPH
def add_table_to_doc(doc, df, heading, table_style='Table Grid'):
    """ Adds a table to a docx document """
    doc.add_heading(heading, level=1).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    columns = list(df.columns)
    # add table
    table = doc.add_table(rows=1, cols=len(columns), style=table_style)
    table.autofit = True
    # add columns (if there is '_' then replace with space)
    for col in range(len(columns)):
        #set_cell_margins(table.cell(0, col), top=100, start=100, bottom=100, end=50)
        table.cell(0, col).text = columns[col].replace("_", " ").capitalize()
    # add data
    for i, row in enumerate(df.itertuples()):
        table_row = table.add_row().cells
        for col in range(len(columns)):
            #set_cell_margins(table_row[col], top=100, start=100, bottom=100, end=50)
            table_row[col].text = str(row[col+1])
    
    return doc



from docx import Document
from docx.shared import Inches
from docx.shared import Mm
import pandas as pd

#Leer archivo
hr_df = pd.read_excel('datos.xlsx')
#Seleccionar columnas
hr_df = hr_df[['A','B']]

#Crear Documento
doc = Document()
#section = doc.section [0]
#section.left_margin = Mm(5)
#section.right_margin = Mm(5)

add_table_to_doc(doc, hr_df.iloc[:7],'')

doc.add_heading('Ductos', level=1)

doc.add_picture('prueba.jpg', width=Inches(1.25))


doc.add_page_break()
doc.save('maqueta.docx')